from abc import ABC, abstractmethod
import docx
import os


class Product:

    def setterDisplay(self, display):
        self._display = display

    def setterCap(self, audiocpu):
        self._audiocpu = audiocpu

    def setterBat(self, battery):
        self._battery = battery

    def setterRam(self, ram):
        self._ram = ram

    def setterCodec(self, codec):
        self._codec = codec


class Player(Product):
    pass


class Document(Product):

    def __init__(self, filename='manual.docx'):
        self.filename = filename

    def create_file(self):
        doc = docx.Document()
        doc.save(self.filename)

    def is_exist_file(self):
        return os.path.exists(self.filename)

    def open_file(self):
        if not self.is_exist_file():
            self.create_file()
        return docx.Document(self.filename)

    def input_info(self):
        doc = self.open_file()
        doc.add_heading("Manual For Player")
        for k, v in self.__dict__.items():
            doc.add_paragraph("{0}:{1}".format(k[1:], str(v)))
        doc.save(self.filename)


class Builder(ABC):
    @abstractmethod
    def setDisplay(self):
        pass

    @abstractmethod
    def setAudioCPU(self):
        pass

    @abstractmethod
    def setBattery(self):
        pass

    @abstractmethod
    def setRam(self):
        pass

    @abstractmethod
    def setCodec(self):
        pass


class PlayerBuilder(Builder):
    def __init__(self):
        self.player = Player()

    def setAudioCPU(self, audiocpu):
        self.player.setterDisplay(audiocpu)

    def setDisplay(self, display):
        self.player.setterDisplay(display)

    def setBattery(self, battery):
        self.player.setterBat(battery)

    def setRam(self, ram):
        self.player.setterRam(ram)

    def setCodec(self, codec):
        self.player.setterCodec(codec)

    def getResult(self):
        return self.player


class ManualBuilder(Builder):
    def __init__(self):
        self.document = Document()

    def setAudioCPU(self, audiocpu):
        self.document.setterCap(audiocpu)

    def setDisplay(self, display):
        self.document.setterDisplay(display)

    def setBattery(self, battery):
        self.document.setterBat(battery)

    def setRam(self, ram):
        self.document.setterRam(ram)

    def setCodec(self, codec):
        self.document.setterCodec(codec)

    def getResult(self):
        self.document.create_file()
        self.document.input_info()
        return self.document


class Director:
    def __init__(self):
        self._builder = None

    def setBuilder(self, builder):
        self._builder = builder

    def getBuilder(self):
        return self._builder

    def buildPlayerCodec(self):
        self._builder.setDisplay("6 дюймов")
        self._builder.setAudioCPU("TFA9003")
        self._builder.setCodec("K-lite")
        self._builder.setBattery("3000Mah")
        self._builder.setRam("6gb")


if __name__ == "__main__":
    builder = PlayerBuilder()
    director = Director()
    director.setBuilder(builder)
    director.buildPlayerCodec()
    # print(director.getBuilder().getResult().__dict__.items())

    builder = ManualBuilder()
    director.setBuilder(builder)
    director.buildPlayerCodec()
    director.getBuilder().getResult()